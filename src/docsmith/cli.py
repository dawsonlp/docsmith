"""YAML-to-Word document generator.

Reads a YAML document file (or stdin) and renders it as a professionally
formatted Word (.docx) document. No computation, no external data -- the
YAML IS the complete document.

Usage:
    docsmith input.yaml
    docsmith input.yaml -o output/
    cat input.yaml | docsmith -
"""

import argparse
import re
import sys
from datetime import UTC, datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from docsmith.normalize import normalize_heading_block, normalize_image_block

APP_NAME = "docsmith"

_alignment_map = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def add_formatted_text(paragraph, text):
    """Add text with inline bold (**) and italic (*) support."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", str(text))
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


def _set_document_properties(doc):
    """Override default metadata set by python-docx's template."""
    now = datetime.now(UTC)
    props = doc.core_properties
    props.author = APP_NAME
    props.comments = "Created by docsmith - https://pypi.org/project/docsmith/"
    props.created = now
    props.modified = now


def _set_compatibility_mode(doc):
    """Set Word compatibility mode to Word 2016+ (version 15)."""
    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = doc.settings.element.makeelement(qn("w:compat"), {})
        settings.append(compat)

    for existing in compat.findall(qn("w:compatSetting")):
        if existing.get(qn("w:name")) == "compatibilityMode":
            compat.remove(existing)

    compat_setting = compat.makeelement(
        qn("w:compatSetting"),
        {
            qn("w:name"): "compatibilityMode",
            qn("w:uri"): "http://schemas.microsoft.com/office/word",
            qn("w:val"): "15",
        },
    )
    compat.append(compat_setting)


def render(doc_data, output_path, base_path=None):
    """Render parsed YAML document data to Word.

    Args:
        doc_data: Dictionary from parsed YAML.
        output_path: Path to write the .docx file.
        base_path: Base directory for resolving relative image paths.
            Required when the document contains image blocks.
    """
    doc = Document()

    title = doc_data.get("title", "")
    subtitle = doc_data.get("subtitle", "")
    status = doc_data.get("status", "")

    _set_document_properties(doc)
    _set_compatibility_mode(doc)

    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(12)

    if status:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Status: {status}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    if title or subtitle or status:
        doc.add_paragraph()

    for block in doc_data.get("content", []):
        if "heading" in block:
            text, level = normalize_heading_block(block)
            doc.add_heading(text, level=level)

        elif "text" in block:
            p = doc.add_paragraph()
            add_formatted_text(p, block["text"])

        elif "bullets" in block:
            for item in block["bullets"]:
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_text(p, item)

        elif "numbered" in block:
            for item in block["numbered"]:
                p = doc.add_paragraph(style="List Number")
                add_formatted_text(p, item)

        elif "table" in block:
            tbl = block["table"]
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, header in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = str(header)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                for row_idx, row_data in enumerate(rows):
                    for col_idx in range(min(len(row_data), len(headers))):
                        cell = table.rows[row_idx + 1].cells[col_idx]
                        cell.text = str(row_data[col_idx])
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                doc.add_paragraph()

        elif "decision" in block:
            p = doc.add_paragraph()
            run = p.add_run("Decision required: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            run.font.size = Pt(11)
            p.add_run(str(block["decision"]))

        elif "image" in block:
            if base_path is None:
                raise ValueError(
                    "Image blocks require a base path for resolving relative "
                    "image paths. This is a programming error -- base_path "
                    "should be passed to render()."
                )
            img = normalize_image_block(block, base_path)
            width = Inches(img.width) if img.width is not None else Inches(5.0)
            doc.add_picture(str(img.path), width=width)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = _alignment_map[img.alignment]
            if img.caption:
                caption_para = doc.add_paragraph()
                caption_para.alignment = last_paragraph.alignment
                run = caption_para.add_run(img.caption)
                run.italic = True
                run.font.size = Pt(9)

    doc.save(str(output_path))


HELP_EPILOG = """\
examples:
  docsmith input.yaml                   # file to Word
  docsmith input.yaml -o reports/       # file to Word in output dir
  cat input.yaml | docsmith -           # pipe YAML from stdin
  llm "write a report" | docsmith -     # pipe from an LLM

YAML format:
  title:    "Document Title"        # optional metadata
  subtitle: "Subtitle"
  status:   "Draft"
  content:                          # list of blocks:
    - heading: "Section"            #   heading at level 1 (default)
    - heading: "Section"            #   heading at level 2
      level: 2                      #     (level is optional, 1-4)
    - heading:                      #   dict form (same result):
        text: "Subsection"
        level: 3
    - text: "Paragraph **bold**"    #   paragraph with **bold**/*italic*
    - bullets: ["Item 1","Item 2"]  #   unordered list
    - numbered: ["Step 1","Step 2"] #   ordered list
    - table:                        #   table with headers + rows
        headers: ["A", "B"]
        rows: [["1","2"]]
    - decision: "Needs review"      #   red decision callout
    - image:                        #   embedded image (PNG/JPEG)
        path: "diagram.png"         #     file path (relative to YAML file)
        width: 5.0                  #     optional width in inches
        alignment: center           #     optional: left, center, right
        caption: "Figure 1"         #     optional caption text

Full docs: https://pypi.org/project/docsmith/
Source:    https://github.com/dawsonlp/docsmith
"""


def main():
    parser = argparse.ArgumentParser(
        description="docsmith -- YAML-to-Word document generator for automated pipelines and LLMs",
        epilog=HELP_EPILOG,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "input",
        help='Path to YAML document file, or "-" to read from stdin',
    )
    parser.add_argument(
        "--output",
        "-o",
        help="Output directory (default: same as input, or cwd for stdin)",
        default=None,
    )
    args = parser.parse_args()

    # Read YAML from file or stdin
    if args.input == "-":
        raw = sys.stdin.read()
        if not raw.strip():
            print("Error: no input received on stdin", file=sys.stderr)
            sys.exit(1)
        doc_data = yaml.safe_load(raw)
        default_output_dir = Path.cwd()
        base_path = Path.cwd()
        output_name = "docsmith_output.docx"
    else:
        input_path = Path(args.input).resolve()
        if not input_path.exists():
            print(f"Error: {input_path} not found", file=sys.stderr)
            sys.exit(1)
        with open(input_path, encoding="utf-8") as f:
            doc_data = yaml.safe_load(f)
        default_output_dir = input_path.parent
        base_path = input_path.parent
        output_name = (
            input_path.stem.replace("_", " ").title().replace(" ", "_") + ".docx"
        )

    if args.output:
        output_dir = Path(args.output).resolve()
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = default_output_dir

    output_path = output_dir / output_name

    render(doc_data, output_path, base_path=base_path)
    print(f"Generated: {output_path}")


if __name__ == "__main__":
    main()
