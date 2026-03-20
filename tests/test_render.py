"""Integration tests for docsmith rendering.

These tests call render() and inspect the generated .docx files.
Domain-level normalization tests are in test_normalize.py.
"""

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from docsmith.cli import render

# ---------------------------------------------------------------------------
# Test image fixtures (generated via Pillow for python-docx compatibility)
# ---------------------------------------------------------------------------


def _make_image_bytes(fmt):
    """Generate minimal valid image bytes using Pillow."""
    import io

    from PIL import Image as PILImage

    img = PILImage.new("RGB", (1, 1), color="red")
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()


@pytest.fixture()
def png_file(tmp_path):
    p = tmp_path / "test.png"
    p.write_bytes(_make_image_bytes("PNG"))
    return p


@pytest.fixture()
def jpeg_file(tmp_path):
    p = tmp_path / "test.jpg"
    p.write_bytes(_make_image_bytes("JPEG"))
    return p


# ---------------------------------------------------------------------------
# Integration tests: heading rendering (kept from original)
# ---------------------------------------------------------------------------


def test_render_flat_heading(tmp_path):
    """Flat form heading renders correct text in generated document."""
    doc_data = {
        "content": [
            {"heading": "Flat Section", "level": 2},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path)

    doc = Document(output_path)
    heading_paragraphs = [
        p for p in doc.paragraphs if p.style.name.startswith("Heading")
    ]
    assert len(heading_paragraphs) == 1
    assert heading_paragraphs[0].text == "Flat Section"


def test_render_nested_heading(tmp_path):
    """Nested dict form heading renders correct text in generated document."""
    doc_data = {
        "content": [
            {"heading": {"text": "Nested Section", "level": 3}},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path)

    doc = Document(output_path)
    heading_paragraphs = [
        p for p in doc.paragraphs if p.style.name.startswith("Heading")
    ]
    assert len(heading_paragraphs) == 1
    assert heading_paragraphs[0].text == "Nested Section"


def test_render_invalid_heading_does_not_produce_file(tmp_path):
    """Invalid heading raises ValueError before producing output file."""
    doc_data = {
        "content": [
            {"heading": {"level": 2}},
        ],
    }
    output_path = tmp_path / "output.docx"
    with pytest.raises(ValueError, match="requires a 'text' key"):
        render(doc_data, output_path)
    assert not output_path.exists()


# ---------------------------------------------------------------------------
# Existing smoke test (must remain unchanged)
# ---------------------------------------------------------------------------


def test_render_minimal(tmp_path):
    """Rendering a minimal document produces a .docx file."""
    doc_data = {
        "title": "Test Document",
        "content": [
            {"text": "Hello, world."},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path)
    assert output_path.exists()
    assert output_path.stat().st_size > 0


# ---------------------------------------------------------------------------
# Integration tests: image rendering
# ---------------------------------------------------------------------------


def test_render_png_image(tmp_path, png_file):
    """PNG image renders in the generated document."""
    doc_data = {
        "content": [
            {"image": {"path": "test.png"}},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path, base_path=tmp_path)

    doc = Document(output_path)
    assert len(doc.inline_shapes) == 1


def test_render_jpeg_image(tmp_path, jpeg_file):
    """JPEG image renders in the generated document."""
    doc_data = {
        "content": [
            {"image": {"path": "test.jpg"}},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path, base_path=tmp_path)

    doc = Document(output_path)
    assert len(doc.inline_shapes) == 1


def test_render_image_with_caption(tmp_path, png_file):
    """Image with caption produces caption paragraph with italic text."""
    doc_data = {
        "content": [
            {"image": {"path": "test.png", "caption": "Figure 1: Architecture"}},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path, base_path=tmp_path)

    doc = Document(output_path)
    # Find caption paragraph -- it follows the image paragraph
    paragraphs = doc.paragraphs
    caption_para = paragraphs[-1]
    assert caption_para.text == "Figure 1: Architecture"
    assert caption_para.runs[0].italic is True
    assert caption_para.runs[0].font.size == Pt(9)


def test_render_image_center_alignment(tmp_path, png_file):
    """Image with center alignment sets paragraph alignment."""
    doc_data = {
        "content": [
            {"image": {"path": "test.png", "alignment": "center"}},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path, base_path=tmp_path)

    doc = Document(output_path)
    # The image paragraph is the last one (no caption)
    image_para = doc.paragraphs[-1]
    assert image_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_render_image_with_width(tmp_path, png_file):
    """Image with explicit width sets inline shape width."""
    doc_data = {
        "content": [
            {"image": {"path": "test.png", "width": 3.0}},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path, base_path=tmp_path)

    doc = Document(output_path)
    shape = doc.inline_shapes[0]
    # Compare with tolerance (EMU rounding for tiny source images)
    assert abs(shape.width - Inches(3.0)) < Inches(0.05)


def test_render_image_default_width(tmp_path, png_file):
    """Image without width gets default width of 5 inches."""
    doc_data = {
        "content": [
            {"image": {"path": "test.png"}},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path, base_path=tmp_path)

    doc = Document(output_path)
    shape = doc.inline_shapes[0]
    assert abs(shape.width - Inches(5.0)) < Inches(0.01)


def test_render_image_default_alignment_is_left(tmp_path, png_file):
    """Image without alignment defaults to left."""
    doc_data = {
        "content": [
            {"image": {"path": "test.png"}},
        ],
    }
    output_path = tmp_path / "output.docx"
    render(doc_data, output_path, base_path=tmp_path)

    doc = Document(output_path)
    image_para = doc.paragraphs[-1]
    # Left alignment may be None (Word default) or explicit LEFT
    assert image_para.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)


def test_render_missing_image_does_not_produce_file(tmp_path):
    """Missing image raises ValueError before producing output file."""
    doc_data = {
        "content": [
            {"image": {"path": "nonexistent.png"}},
        ],
    }
    output_path = tmp_path / "output.docx"
    with pytest.raises(ValueError, match="not found"):
        render(doc_data, output_path, base_path=tmp_path)
    assert not output_path.exists()
